"""DOCX and TXT document processing."""

from __future__ import annotations

import io
from typing import Any, Dict

from utils.text_extractor import extract_code_from_text, extract_errors_from_text


def process_document(
    file_path: str,
    filename: str,
    ext: str,
) -> Dict[str, Any]:
    """Extract text from DOCX or TXT files.

    Args:
        file_path: Path to the temporary file.
        filename: Original name.
        ext: Lowercase extension including dot.

    Returns:
        Dict with text, code, error_message, success.
    """
    text = ""
    ext = ext.lower()

    if ext == ".txt":
        try:
            with open(file_path, "rb") as f:
                raw_bytes = f.read()

            for enc in ("utf-8", "latin-1", "cp1252"):
                try:
                    text = raw_bytes.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue

            if not text.strip() and len(raw_bytes) > 0:
                text = raw_bytes.decode("utf-8", errors="replace")
        except Exception as exc:
            return {
                "text": "",
                "code": "",
                "error_message": "",
                "success": False,
                "error": f"TXT read failed: {exc}",
            }
    elif ext == ".docx":
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = []
            for p in doc.paragraphs:
                paragraphs.append(p.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            text = "\n".join(paragraphs).strip()
        except Exception as exc:
            return {
                "text": "",
                "code": "",
                "error_message": "",
                "success": False,
                "error": f"DOCX read failed: {exc}. Install python-docx.",
            }
    else:
        return {
            "text": "",
            "code": "",
            "success": False,
            "error": f"Unsupported document type: {ext}",
        }

    if not text.strip():
        return {
            "text": "",
            "code": "",
            "success": False,
            "error": f"Document '{filename}' is empty.",
        }

    return {
        "text": text,
        "code": extract_code_from_text(text),
        "error_message": extract_errors_from_text(text),
        "transcript": "",
        "success": True,
        "error": None,
        "source": "document",
        "filename": filename,
    }
